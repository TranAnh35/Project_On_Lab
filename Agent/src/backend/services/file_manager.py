import os
from fastapi import UploadFile
import PyPDF2
from docx import Document
import yaml
from typing import List

UPLOAD_FOLDER = "uploaded_files"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def save_file(file: UploadFile) -> str:
    file_path = os.path.join(UPLOAD_FOLDER, file.filename)
    with open(file_path, "wb") as f:
        f.write(file.file.read())
    return file_path

def delete_file(file_name: str) -> bool:
    file_path = os.path.join(UPLOAD_FOLDER, file_name)
    if os.path.exists(file_path):
        os.remove(file_path)
        return True
    return False

def list_files() -> List[dict]:
    files = []
    for file_name in os.listdir(UPLOAD_FOLDER):
        file_path = os.path.join(UPLOAD_FOLDER, file_name)
        size = os.path.getsize(file_path)
        files.append({"name": file_name, "size": size, "path": file_path})
    return files

def read_file(file_path: str) -> str:
    """Đọc nội dung file dựa trên định dạng."""
    file_name = os.path.basename(file_path)
    if file_name.endswith('.pdf'):
        return read_pdf(file_path)
    elif file_name.endswith(('.doc', '.docx')):
        return read_docx(file_path)
    elif file_name.endswith(('.yaml', '.yml')):
        return read_yaml(file_path)
    else:
        return read_txt_file(file_path)

def read_pdf(file_path: str) -> str:
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text

def read_docx(file_path: str) -> str:
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def read_yaml(file_path: str) -> str:
    with open(file_path, 'r', encoding='utf-8') as f:
        data = yaml.safe_load(f)
        return str(data)

def read_txt_file(file_path: str) -> str:
    encodings = ['utf-8', 'latin1', 'cp1252', 'iso-8859-1']
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    raise Exception(f"Không thể đọc file với các encoding đã thử")

async def read_uploaded_file(file: UploadFile) -> str:
    """Đọc nội dung file được gửi trực tiếp từ frontend dựa trên định dạng."""
    file_name = file.filename
    content: str = ""
    
    try:
        if file_name.endswith('.pdf'):
            content = await read_pdf(file)
        elif file_name.endswith(('.doc', '.docx')):
            content = await read_docx(file)
        elif file_name.endswith(('.yaml', '.yml')):
            content = await read_yaml(file)
        else:
            content = await read_txt_file(file)
        return content
    except Exception as e:
        raise Exception(f"Không thể đọc file {file_name}: {str(e)}")
    finally:
        await file.close()  # Đóng file sau khi đọc

async def read_pdf(file: UploadFile) -> str:
    """Đọc nội dung file PDF từ UploadFile."""
    pdf_reader = PyPDF2.PdfReader(file.file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text

async def read_docx(file: UploadFile) -> str:
    """Đọc nội dung file DOCX từ UploadFile."""
    doc = Document(file.file)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

async def read_yaml(file: UploadFile) -> str:
    """Đọc nội dung file YAML từ UploadFile."""
    content = await file.read()
    data = yaml.safe_load(content)
    return str(data)

async def read_txt_file(file: UploadFile) -> str:
    """Đọc nội dung file text từ UploadFile."""
    content = await file.read()
    encodings = ['utf-8', 'latin1', 'cp1252', 'iso-8859-1']
    for encoding in encodings:
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise Exception(f"Không thể đọc file với các encoding đã thử")